import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    """Sets background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets padding for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_styled_document():
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set default style font
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(10.5)
    normal_font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    
    return doc

def parse_markdown_to_docx(md_path, docx_path, doc_title):
    doc = create_styled_document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_buffer = []
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\r\n')
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                in_code_block = False
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(6)
                code_text = "\n".join(code_buffer)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                code_buffer = []
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue
            
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
            
        # Handle Markdown Tables
        if '|' in line and not line.strip().startswith('```'):
            if not in_table:
                in_table = True
                table_rows = []
            if re.match(r'^\s*\|?\s*[-:]+[-| :]*\|\s*$', line) or re.match(r'^[\s|:-]+$', line):
                i += 1
                continue
            cols = [c.strip() for c in line.split('|')]
            if cols and cols[0] == '':
                cols = cols[1:]
            if cols and cols[-1] == '':
                cols = cols[:-1]
            if cols:
                table_rows.append(cols)
            i += 1
            continue
        else:
            if in_table:
                if len(table_rows) > 0:
                    num_cols = max(len(r) for r in table_rows)
                    tbl = doc.add_table(rows=len(table_rows), cols=num_cols)
                    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    for r_idx, row_data in enumerate(table_rows):
                        row = tbl.rows[r_idx]
                        is_header = (r_idx == 0)
                        for c_idx in range(num_cols):
                            cell = row.cells[c_idx]
                            cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
                            clean_text = cell_text.replace('**', '').replace('`', '')
                            cell.text = clean_text
                            
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_before = Pt(3)
                            p.paragraph_format.space_after = Pt(3)
                            if len(p.runs) > 0:
                                p.runs[0].font.name = 'Calibri'
                                p.runs[0].font.size = Pt(9.5)
                                if is_header:
                                    p.runs[0].font.bold = True
                                    p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                else:
                                    p.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                            
                            if is_header:
                                set_cell_background(cell, "1E3A5F")
                            else:
                                if r_idx % 2 == 1:
                                    set_cell_background(cell, "F8FAFC")
                                else:
                                    set_cell_background(cell, "FFFFFF")
                            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                            
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)
                in_table = False
                table_rows = []
                
        if not line.strip():
            i += 1
            continue
            
        if line.startswith('# '):
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0F, 0x2A, 0x4A)
            i += 1
            continue
        elif line.startswith('## '):
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[3:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            i += 1
            continue
        elif line.startswith('### '):
            p = doc.add_heading(level=3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(line[4:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2B, 0x4C, 0x7E)
            i += 1
            continue
        elif line.startswith('#### '):
            p = doc.add_heading(level=4)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[5:].strip())
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x3B, 0x59, 0x88)
            i += 1
            continue
            
        if line.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            clean_line = line.lstrip('> ').strip()
            add_formatted_runs(p, clean_line, italic_default=True, color=RGBColor(0x1E, 0x3A, 0x5F))
            i += 1
            continue
            
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            indent_level = (len(line) - len(line.lstrip())) // 2
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            clean_line = re.sub(r'^\s*[-*]\s+', '', line)
            add_formatted_runs(p, clean_line)
            i += 1
            continue
            
        num_match = re.match(r'^\s*(\d+)\.\s+(.*)$', line)
        if num_match:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            clean_line = num_match.group(2)
            add_formatted_runs(p, clean_line)
            i += 1
            continue
            
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        add_formatted_runs(p, line)
        i += 1

    os.makedirs(os.path.dirname(docx_path) or '.', exist_ok=True)
    try:
        doc.save(docx_path)
        print(f"Generated DOCX: {docx_path}")
    except PermissionError:
        alt_path = docx_path.replace('.docx', '_new.docx')
        doc.save(alt_path)
        print(f"Generated DOCX (Alternate): {alt_path} (Original file was locked)")

def add_formatted_runs(paragraph, text, italic_default=False, color=None):
    """Parses markdown bold, italic, code, and math symbols into Word runs."""
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    tokens = pattern.split(text)
    
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            if color:
                run.font.color.rgb = color
            else:
                run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
        elif token.startswith('`') and token.endswith('`'):
            run = paragraph.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            if color:
                run.font.color.rgb = color
        else:
            run = paragraph.add_run(token)
            run.font.name = 'Calibri'
            run.font.size = Pt(10.5)
            if italic_default:
                run.italic = True
            if color:
                run.font.color.rgb = color

if __name__ == '__main__':
    docs_to_convert = [
        ('README.md', 'README.docx', 'Project Documentation & Technical Guide'),
        ('docs/DEMO_SCRIPT.md', 'docs/DEMO_SCRIPT.docx', '3-Minute Live Hackathon Pitch Script'),
        ('docs/JUDGE_QA.md', 'docs/JUDGE_QA.docx', 'Anticipated Judge Q&A and Defense Strategies'),
        ('docs/architecture.md', 'docs/ARCHITECTURE.docx', 'System Architecture & Data Flow Specification'),
        ('docs/PROJECT_ANALYSIS_REPORT.md', 'docs/PROJECT_ANALYSIS_REPORT.docx', 'Comprehensive Project Analysis Report'),
    ]
    
    for md_file, docx_file, title in docs_to_convert:
        if os.path.exists(md_file):
            parse_markdown_to_docx(md_file, docx_file, title)
        else:
            print(f"File not found: {md_file}")
